"""
File utilities for extracting text from various file formats.
"""
import io
from typing import Optional


def extract_text_from_file(content: bytes, filename: str) -> str:
    """
    Extract text from file content based on filename extension.
    
    Args:
        content: Raw file bytes
        filename: Original filename with extension
        
    Returns:
        Extracted text content or empty string if extraction fails
    """
    if not filename:
        return _try_utf8_decode(content)
    
    filename_lower = filename.lower()
    
    # PDF files
    if filename_lower.endswith('.pdf'):
        return _extract_pdf_text(content)
    
    # DOCX files
    elif filename_lower.endswith('.docx'):
        return _extract_docx_text(content)
    
    # Plain text files (fallback for everything else)
    else:
        return _try_utf8_decode(content)


def _try_utf8_decode(content: bytes) -> str:
    """Try to decode bytes as UTF-8, return empty string on failure."""
    try:
        return content.decode('utf-8')
    except UnicodeDecodeError:
        # Try other common encodings
        for encoding in ['latin1', 'cp1252', 'iso-8859-1']:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        return ""


def _extract_pdf_text(content: bytes) -> str:
    """Extract text from PDF content."""
    try:
        from pypdf import PdfReader
        
        pdf_file = io.BytesIO(content)
        reader = PdfReader(pdf_file)
        
        text_parts = []
        for page in reader.pages:
            try:
                text_parts.append(page.extract_text())
            except Exception:
                # Skip pages that can't be extracted
                continue
        
        return '\n'.join(text_parts)
    
    except ImportError:
        # pypdf not available, try fallback
        return ""
    except Exception:
        # PDF parsing failed
        return ""


def _extract_docx_text(content: bytes) -> str:
    """Extract text from DOCX content."""
    try:
        from docx import Document
        
        docx_file = io.BytesIO(content)
        doc = Document(docx_file)
        
        text_parts = []
        for paragraph in doc.paragraphs:
            text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        
        return '\n'.join(text_parts)
    
    except ImportError:
        # python-docx not available
        return ""
    except Exception:
        # DOCX parsing failed
        return ""